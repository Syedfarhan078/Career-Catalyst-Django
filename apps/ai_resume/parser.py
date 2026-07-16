import re
import pdfplumber
import docx

def clean_text(text):
    """
    Remove non-printable characters and clean extra whitespaces.
    """
    if not text:
        return ""
    # Replace multiple whitespaces/newlines with single spaces/newlines
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_text_from_pdf(file):
    """
    Extract raw text from a PDF file using pdfplumber.
    """
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        raise ValueError("Could not parse PDF file. Ensure it is a valid, text-based PDF.")
    return clean_text(text)

def extract_text_from_docx(file):
    """
    Extract raw text from a DOCX file using python-docx.
    """
    text = []
    try:
        doc = docx.Document(file)
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        # Also extract text from tables inside docx
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        raise ValueError("Could not parse DOCX file. Ensure it is a valid Word document.")
    return clean_text("\n".join(text))

def parse_resume_file(file):
    """
    Determine file type (PDF/DOCX) and extract text.
    """
    filename = file.name.lower()
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        return extract_text_from_docx(file)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
